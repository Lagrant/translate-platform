import ast
import copy
import csv
import os
import re
from io import StringIO
import sys
from .util import ListConverter
from .pdf_operator import merge, merge_all, split
import platform
import random
import shutil
import smtplib
from datetime import timedelta, datetime
from email.header import Header
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import pymysql
from flask import Flask, request, json, render_template, session, jsonify, url_for, current_app, g, redirect, Response, flash, send_file, make_response, send_from_directory
from xlrd import open_workbook
from .model import user, login, mailconfirm, db, book, task, project, translators_tasks
from werkzeug.utils import secure_filename
from PyPDF2 import PdfFileReader, PdfFileWriter
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.pagesizes import letter
import reportlab.pdfbase.ttfonts
reportlab.pdfbase.pdfmetrics.registerFont(reportlab.pdfbase.ttfonts.TTFont('SimSun', 'SimSun.ttf'))
from . import app
from .SQLEncoder import AlchemyEncoder
from .miner import parse
from .ahocorasick import AhoCorasick
from .fanyigou import upload_to_fanyigou, query_process, download_file

log = app.logger

@app.route('/')
@app.route('/index')
def index():
    g.count = 0
    
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        print(user1)
        return render_template('regulator_homepage.html', user=user1)
    else:
        return render_template('regulator_homepage.html')


@app.route('/login/', methods=['GET', 'POST'])
def user_login():
    return render_template('user/login.html')


# 验证密码
@app.route('/login/pass/', methods=['GET', 'POST'])
def login_pass():
    # 添加数据到session中
    data = request.get_json('data')
    email = data['email']
    pas = data['password']
    theuser = user.query.filter_by(email=email).first()
    if theuser is None:
        return "account not exist"
    elif not theuser.check_password_hash(pas):
        return "password not right"
    else:
        session['email'] = email
        session['user_id'] = theuser.id
        session.permanent = True
        login1 = login(email=email)
        db.session.add(login1)
        db.session.commit()

        if session.get("last_page"):
            print(session.get("last_page"))
            page = session.get("last_page")
            return page
        elif theuser.role == 1:
            return '/'
        else: return '/translator'


@app.route('/login/pass/name/', methods=['GET', 'POST'])
def login_pass_name():
    name = request.get_json()['name']
    user1 = user.query.filter_by(username=name).first()
    if user1 is not None:
        return "true"
    else:
        return "false"


# 发送邮件
def sendmail(to_mail, num):
    # 邮件外主体
    smtp = None
    smtpserver = "smtp.qq.com"
    smtpport = 465
    from_mail = "1361377791@qq.com"
    password = "ejpulrvmshuyibba"
    # 邮件内容主体
    subject = "Activate your account"
    from_name = "Translation"
    body = num + "Here is your verification code, please fill in within five minutes. " \
                 "Ignore this message if it is not my operation.\n"
    msgtext = MIMEText(body, "plain", "utf-8")
    msg = MIMEMultipart()
    msg['Subject'] = Header(subject, "utf-8")
    msg["From"] = Header(from_name + "<" + from_mail + ">", "utf-8")
    msg["To"] = to_mail
    msg.attach(msgtext)
    try:
        smtp = smtplib.SMTP_SSL(smtpserver, smtpport)
        # smtp.starttls()
        smtp.login(from_mail, password)
        smtp.sendmail(from_mail, to_mail, msg.as_string())
        smtp.quit()
        return True
    except Exception as e:
        print(e)
        if smtp is not None:
            smtp.quit()
        return False


def mycopyfile(srcfile, dstfile):
    if not os.path.isfile(srcfile):
        return False
    else:
        fpath, fname = os.path.split(dstfile)
        if not os.path.exists(fpath):
            os.makedirs(fpath)
        shutil.copyfile(srcfile, dstfile)
        print("copy %s->%s" % (srcfile, dstfile))


# 前台传过来的数据都是已经被验证好格式的，传到后台以后只需要进行数据库的比对即可
# 需要验证的内容为：数据库中是否已经存在该用户。如果存在，那么弹出提示信息
# 新建一个用户需要完成的内容有：
# 给该用户的邮箱发送邮件；为该用户在服务器新增一个个人文件夹存储个人信息
# 由于涉及到文件的路径，管理起来有点麻烦所以暂时先不考虑
@app.route('/login/signup/', methods=['GET', 'POST'])
def login_signup():
    data = request.get_json('data')
    email = data['email']
    cur_dir = "./data/users"
    if not os.path.exists('./data/users/' + email):
        os.makedirs('./data/users/' + email)
        cur_dir = cur_dir + "/" + email
    else:
        return "email already exist"
    # if os.path.isdir(cur_dir):
    #     if not os.path.exists('./data/users/' + email):
    #         os.makedirs('./data/users/' + email)
    #         cur_dir = cur_dir + "/" + email
    #     else:
    #         return "email already exist"
        
    # else:
    #     return "error"
    verify = data['verify']
    confirm1 = mailconfirm.search(email, verify)
    if confirm1 is not None:  # 首先看验证码是否正确
        theuser = user.query.filter_by(email=email).first()
        if theuser is not None:  # 然后看用户是否存在
            return "email already exist"
        name = data['username']
        pas = data['password']
        typ = data['tp']  # 以数字的形式存储用户角色
        if typ == 'Manager':
            typ = 1
        elif typ == 'Translator':
            typ = 0
        user1 = user(email=email, username=name, password=pas, role=typ)
        db.session.add(user1)
        db.session.delete(confirm1)
        db.session.commit()
        session['email'] = email
        session.permanent = True
        return "success"
    else:  # 验证码不正确或者已经过期
        return "Verification code error"


# 发送验证码并将验证码存到数据库
@app.route('/login/verify/', methods=['GET', 'POST'])
def login_verify():
    data = request.get_json('data')
    email = data['email']
    num = str(random.randint(1000, 9999))
    send = sendmail(email, num)
    if send:
        confirm1 = mailconfirm.query.filter_by(email=email).first()
        if confirm1 is not None:
            confirm1.num = num
            confirm1.invalid = datetime.now() + timedelta(minutes=5)
        else:
            conf1 = mailconfirm(email=email, num=num)
            db.session.add(conf1)
        db.session.commit()
        return "success"
    else:
        return "fail to send the mail"


# 验证验证码，过期或者成功都删除这个验证码
@app.route('/forget/verify/', methods=['GET', 'POST'])
def forget_verify():
    data = request.get_json('data')
    email = data['email']
    verify = data['verify']
    confirm1 = mailconfirm.query.filter_by(email=email, num=verify).first()
    if confirm1 is not None:
        db.session.delete(confirm1)
        db.session.commit()
        if confirm1.invalid > datetime.now():
            print("session deleted.")
            return "success"
        else:
            return "false"
    else:
        return "false"


@app.route('/forget/change/', methods=['GET', 'POST'])
def forget_change():
    data = request.get_json('data')
    email = data['email']
    password = data['password']
    print("email")
    user1 = user.query.filter_by(email=email).first()
    if user1 is not None:
        user1.password(password)
        db.session.commit()
        return "success"
    else:
        return "false"


@app.route('/user/')
def user_user():
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1 is not None:
            email = user1.email
            cur_dir = './data/users/' + email
            print(cur_dir)
            if os.path.exists(cur_dir):
                return render_template('user/user.html', user=user1)
            else:
                return "Unknown error"

        else:
            return "404 NOT FOUND"
    return render_template('user/login.html')


# 下一步，需要将用户的头像进行压缩
@app.route('/user/change/', methods=['GET', 'POST'])
def user_change():
    email = request.form.get('email')
    name = request.form.get('username')
    password = request.form.get('password')
    gender = request.form.get('gender')
    signature = request.form.get('signature')
    user1 = user.query.filter_by(email=email).first()
    if user1 is not None:
        user1.username = name
        user1.password(password)
        if gender == 'male':
            user1.gender = True
        else:
            user1.gender = False
        user1.signature = signature
        db.session.commit()
        return "success"
    else:
        return "false"


@app.route('/user/change/img/', methods=['GET', 'POST'])
def user_change_img():
    file = request.files['file']
    if file and '.' in file.filename:
        file_types = ['jpg','jpeg','png','pdf']
        this_type = file.filename.rsplit('.', 1)[1]
        for file_type in file_types:
            old_file = './data/users/' + session.get('email') + '/img/user_img.'+file_type
            if os.path.exists(old_file):
                os.remove(old_file)
        if this_type in file_types:
            old_file = './data/users/' + session.get('email') + '/img/user_img.jpg'
            file.save(old_file)
            return 'success'
    else:
        return "filename invalid or network error"


@app.route('/products', methods=['POST', 'GET'])
def products():
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role != 1:
            return "false"
        return render_template('products.html', user=user1)
    else:
        return render_template('products.html')


@app.route('/contact_us', methods=['POST', 'GET'])
def contact_us():
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role != 1:
            return "false"
        return render_template('contact_us.html', user=user1)
    else:
        return render_template('contact_us.html')


@app.route('/about_us', methods=['POST', 'GET'])
def about_us():
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role != 1:
            return "false"
        return render_template('about_us.html', user=user1)
    else:
        return render_template('about_us.html')


@app.route('/whether_login', methods=['POST','GET'])
def whether_login():
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role != 1:
            return "false"
        return "true"
    return "false"

@app.route('/upload_file/<filename>', methods=['POST','GET'])
def upload_file(filename):
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role != 1:
            return "false"
        if request.method == 'POST':
            if 'file' not in request.files:
                flash('No File')
                return redirect(request.url)
            
            project_id = request.form.get("projectId", type = str, default = None)
            current_project = project.query.filter_by(id=project_id).first()
            project_name = current_project.name
            cur_dir = "./data/users/" + email + "/" + project_name
            
            f = request.files['file']
            fname = secure_filename(f.filename)
            path_file_name = cur_dir + "/" + fname
            f.save(path_file_name)

            file_type = fname.split(".")[-1]
            language = current_project.from_language
            
            try:
                parse(path_file_name)
            except Exception as e:
                log_path = './error/'
                if not os.path.exists(log_path):
                    os.makedirs(log_path)
                with open(os.path.join(log_path, 'pdfminer.log'), 'a') as f:
                    f.write(e)
                return "fail to convert pdf to txt file"
            
            pdf_reader = PdfFileReader(path_file_name)
            page_num = pdf_reader.getNumPages()
            
            book1 = book(book_name = path_file_name, file_type = file_type, language = language, page_number = page_num, project_id = project_id)
            db.session.add(book1)
            db.session.commit()
            return Response('Uploaded file successfully', status=200)
        else:
            return "false"
    else:
        return "false"

@app.route('/page_number/<projectID>', methods=['POST','GET'])
def get_page_num(projectID):
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role != 1:
            return "false"
        if request.method == 'POST':
            current_book = book.query.filter_by(project_id=projectID).first()
            page_number = current_book.translated_page_number
            return jsonify(page_number)
    else:
        return "false" 

@app.route('/set_setting_list', methods=['POST','GET'])
def set_setting_list():
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role != 1:
            return "false"
        if request.method == 'POST':
            setting_list = json.loads(request.get_data())
            cur_dir = "./data/users/" + email
            cur_dir = cur_dir + "/" + setting_list["name"]
            if not os.path.exists(cur_dir):
                os.makedirs(cur_dir)
            project1 = project(name = setting_list["name"], manager_id = user1.id, deadline = setting_list["ddl"],
             href = setting_list["href"], from_language = setting_list["language"])
            db.session.add(project1)
            db.session.commit()

            response_list = {}
            response_list["projectId"] = project1.id
            response_list["status"] = 200
            return jsonify(response_list)
    else:
        return "false"

def compute_progress(proj_id):
    all_tasks = task.query.filter_by(project_id=proj_id).all()
    finished_task = 0
    total_task = 0
    for single_task in all_tasks:
        total_task += translators_tasks.query.filter_by(task_id=single_task.id).count()
        finished_task += translators_tasks.query.filter_by(task_id=single_task.id, task_type="finished").count()
    if total_task == 0:
        rate = 0
    else:
        rate = int(finished_task / total_task * 100)
    progress = str(rate) + "%"
    return progress

@app.route('/get_setting_list', methods=['POST','GET'])
def get_setting_list():
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role is None:
            return "404 NOT FOUND"
        if request.method == 'POST':
            if user1.role == 1:
                # 当user为manager时，从project中查找项目
                item_list = project.query.filter_by(manager_id=user1.id).all()
                item_list_dict = []
                for i_list in item_list:
                    cur_progress = compute_progress(i_list.id)
                    i_list.progress = cur_progress
                    db.session.add(i_list)
                    db.session.commit()
                    item_list_dict.append(json.dumps(i_list, cls = AlchemyEncoder))
                return jsonify(item_list_dict)
            else:
                # 当user为translator时，从task中查找项目
                # item_list = task.query.filter_by(translators=user1).all()
                item_list = user1.tasks
                item_list_dict = []
                for i_list in item_list:
                    cur_task = task.query.filter_by(id=i_list.task_id).first()
                    item = {}
                    i_proj = project.query.filter_by(id=cur_task.project_id).first()
                    item['name'] = i_proj.name
                    item['type'] = i_list.task_type
                    item['deadline'] = i_proj.deadline
                    item['taskId'] = cur_task.id
                    item['progress'] = 0
                    item_list_dict.append(item)
                return jsonify(item_list_dict)
        else:
            return "false"
    else:
        return "false"

@app.route('/get_setting_item/<string:name>', methods=['POST','GET'])
def get_setting_item(name):
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role is None:
            return "false"
        if request.method == 'POST':
            item = project.query.filter_by(name = name).first()
            return json.dumps(item, cls = AlchemyEncoder)
        else:
            return "false"
    else:
        return "false"

@app.route('/send_segment_range/<string:name>', methods=['POST','GET'])
def segment_file(name):
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role != 1:
            return "false"
        if request.method == 'POST':
            select_range = json.loads(request.get_data())

            current_project = project.query.filter_by(name=name).first()
            current_project_id = current_project.id

            current_book = book.query.filter_by(project_id=current_project_id).first()
            current_book_name = current_book.translated_book_name

            try:
                out_files = split(current_book_name, select_range)
            except Exception as e:
                log_path = './error/'
                if not os.path.exists(log_path):
                    os.makedirs(log_path)
                with open(os.path.join(log_path, 'pypdf2.log'), 'a') as f:
                    f.write(e)
                return "fail to segment file"

            for i in range(len(select_range)):
                parse(out_files[i])
                temp = list(map(int, select_range[i]))
                start_page = temp[0] - 1
                end_page = temp[1] - 1

                task1 = task(project_id=current_project_id, start_page=start_page, end_page=end_page, task_path=out_files[i])
                db.session.add(task1)
                db.session.commit()

            return Response("Successfully segment the file", status=200)
        else:
            return "false"
    else:
        return "false"
        
@app.route('/terms/<string:name>', methods=['POST','GET'])
def replace_terms(name):
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role != 1:
            return "false"
        if request.method == 'POST':
            l_d = json.loads(request.get_data())
            target_language = l_d[0]

            current_project = project.query.filter_by(name=name).first()
            current_project.to_language = target_language
            current_project_id = current_project.id
            current_book = book.query.filter_by(project_id=current_project_id).first()
            current_book_name = current_book.book_name
            source_language = current_book.language
            text_current_book_name = current_book_name[0: current_book_name.rindex(".")] + ".txt"
            term_current_book_name = current_book_name[0: current_book_name.rindex(".")] + "_term.txt"
            db.session.add(current_project)
            db.session.commit()

            target_language_index = {}
            source_language_string = []

            if not os.path.exists("./terms/wording-cn-jp.csv"):
                # return Response(data=json.dumps({'false': "term file not found"}), mimetype='application/json', status=404)
                return "term file not found"

            with open("./terms/wording-cn-jp.csv", "r") as f:
                line = f.readline()

                while line:
                    line = line.strip()
                    words = line.split(",")
                    if source_language == "en":
                        source_language_string.append(words[0])
                    elif source_language == "zh":
                        source_language_string.append(words[1])
                    elif source_language == "jp":
                        source_language_string.append(words[2])
                    if target_language == "en":
                        target_language_index[source_language_string[-1]] = words[0]
                    elif target_language == "zh":
                        target_language_index[source_language_string[-1]] = words[1]
                    elif target_language == "jp":
                        target_language_index[source_language_string[-1]] = words[2]
                    
                    line = f.readline()
            
            if not os.path.exists(text_current_book_name):
                return "fail to convert pdf to txt file"
            
            with open(text_current_book_name, "r") as f:
                
                line = f.readline()
                text = ""
                while line:
                    end_character = line[-1]
                    tree=AhoCorasick(*source_language_string)
                    results = tree.search(line, True)

                    lr = []
                    while len(results) > 0:
                        lr.append(results.pop())
                    lr = sorted(lr, key=lambda x: x[1][0], reverse=True)

                    for i in range(0,len(lr)):
                        source_language_index = lr[i][0]
                        start_index = lr[i][1][0]
                        end_index = lr[i][1][1]

                        substring1 = line[0:start_index]
                        substring2 = line[end_index:-1]
                        line = substring1 + target_language_index[source_language_index] + substring2 + end_character
                    text += line
                    line = f.readline()
                
                with open(term_current_book_name, "w", encoding="utf-8") as fl:
                        fl.write(text)
            return "true"
    else:
        return "false"

@app.route('/get_range_list/<string:name>', methods=['POST','GET'])
def get_range_list(name):
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role != 1:
            return "false"
        if request.method == 'POST':
            range_list = []
            current_project = project.query.filter_by(name=name).first()
            specific_tasks = task.query.filter_by(project_id=current_project.id).all()
            for specific_task in specific_tasks:
                range_list_item = [specific_task.start_page + 1, specific_task.end_page + 1]
                range_list.append(range_list_item)
            
            return jsonify(range_list)
        else:
            return "false"
    else:
        return "false"

@app.route('/translator_task/<list:ids>', methods=['POST','GET'])
def translator_setion(ids):
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role != 1:
            return "false"
        if request.method == 'POST':

            name = ids[0]
            assign_type = ids[1]
            t_s = json.loads(request.get_data()) #t_s: translator & task

            sections = list(t_s.values())
            translators = list(t_s.keys())
            current_project = project.query.filter_by(name=name).first()
            for i in range(0,len(t_s)):
                spage = int(sections[i].split(",")[0]) - 1
                epage = int(sections[i].split(",")[1]) - 1
                cur_task = task.query.filter_by(project_id=current_project.id, 
                    start_page=spage, end_page=epage).first()
                if cur_task is None:
                    return 'Task not found ' + sections[i]
                duplicated = translators_tasks.query.filter_by(task_id=cur_task.id, task_type=assign_type)
                if not duplicated is None:
                    return "duplicated task in this assignment"
                translator = user.query.filter_by(username=translators[i]).first()
                if translator is None:
                    return 'Translator not found ' + translators[i]
                if translator.role != 0:
                    return 'Invalide translator role'
                
                # cur_task.translators.append(translator)
                # translator.tasks.append(cur_task)
                translators_tasks1 = translators_tasks(translator.id, cur_task.id, assign_type)
                db.session.add(translators_tasks1)
                db.session.add(cur_task)
                db.session.add(translator)
                db.session.commit()

            return "true"
        else:
            return "false"
    else: 
        return "false"

# @app.route('/get_translator_task_file/<string:name>', methods=['POST','GET'])
# def get_t_s_file(name):
    
#     if request.method == 'POST':
        
        
#         translator = json.loads(request.get_data())
#         t_s = n_t_s[name]
#         section = t_s[translator]
#         filename = out_files_list[section]
#         return jsonify(filename)

@app.route("/files/<path:path>", methods=['POST','GET'])
def send_source_file(path):
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role is None:
            return "false"
        return send_file("../"+path)
    else: 
        return "false"


@app.route('/translation/<list:ids>', methods=['POST','GET'])
def toTranslationPage(ids):
    """
    call translation api
    """
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role is None:
            return "false"
        current_project = project.query.filter_by(name=ids[0]).first()
        current_book = book.query.filter_by(project_id=current_project.id).first()
        projId = current_book.book_name
        manager = current_project.manager
        term_book = projId[0: projId.rindex(".")] + "_term.txt"

        #把*_term.txt文件转为*_term.docx文件，然后上传翻译狗
        # docx_term_book = projId[0: projId.rindex(".")] + "_term.docx"
        # document = Document()
        # myfile = open(term_book, "rb").read()
        # myfile=myfile.decode("utf-8","ignore")
        # #myfile = re.sub(r'[^\x00-\x7F]+|\x0c',' ', myfile) # remove all non-XML-compatible characters
        # p = document.add_paragraph(myfile)
        # document.save(docx_term_book)
        # myfile.close()

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(fontName='SimSun', name='Song', leading=20, fontSize=12))
        styleN = styles['Song']
        story = []
        pdf_term_book = projId[0: projId.rindex(".")] + "_term.pdf"
        # PDFCreator(term_book)
        doc = SimpleDocTemplate(
            pdf_term_book,
            pagesize=letter,
            bottomMargin=.4 * inch,
            topMargin=.6 * inch,
            rightMargin=.8 * inch,
            leftMargin=.8 * inch)
        with open(term_book, "r") as f:
            text_content = f.read()
        P = Paragraph(text_content, styleN)
        story.append(P)

        doc.build(
            story,
        )

        translationId = "./data/users/" + manager.email + "/" + current_project.name + "/" + ids[1]
        current_book.translated_book_name = translationId
        db.session.add(current_book)
        db.session.commit()
        if not os.path.isfile(translationId):
            # 变量要改成 docx_term_book #不用了
            tid, msg = upload_to_fanyigou(current_book_name=pdf_term_book, source_language=current_project.from_language, target_language=current_project.to_language)
            if tid == False:
                return msg
        else: 
            tid = -1
        
        return render_template('translationpage.html', tid = tid, projId = projId, translationId = translationId)
    else:
        return "false"

@app.route('/query/<tid>', methods=['POST','GET'])
def query(tid):
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role != 1:
            return "false"
        if request.method == 'POST':
            #查询翻译进度
            process_data = query_process(tid)
            if process_data == False:
                return "false"
            return jsonify(process_data)
        else:
            return "false"
    else: 
        return "false"    

@app.route('/download_file/<path:path>', methods=['POST','GET'])
def download(path):
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role != 1:
            return "false"
        if request.method == 'POST':
            #下载翻译结果
            # ids[0] = cur_dir, ids[1] = tid
            tid = json.loads(request.get_data())
            if os.path.isfile("./"+path):
                return "true"
            feed_back = download_file("./"+path,tid)
            # feed_back = True
            if feed_back:
                #更新book表单
                current_book = book.query.filter_by(translated_book_name="./"+path).first()
                pdf_reader = PdfFileReader("./"+path)
                page_num = pdf_reader.getNumPages()
                current_book.translated_page_number = page_num
                db.session.add(current_book)
                db.session.commit()
                return "true"
            else: 
                return "false"
        else:
            return "false"
    else: 
        return "false"

@app.route('/get_task_list/<name>', methods=['POST','GET'])
def get_task_list(name):
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role != 1:
            return "false"
        if request.method == 'POST':
            
            task_item_list = []
            cur_proj = project.query.filter_by(name=name).first()
            tasks = task.query.filter_by(project_id=cur_proj.id).all()
            for cur_task in tasks:
                trans_tasks = translators_tasks.query.filter_by(task_id=cur_task.id).all()
                for tran_task in trans_tasks:
                    task_item = {}
                    task_item["translator"] = tran_task.translator.username
                    task_item["status"] = tran_task.task_type
                    task_item["view"] = "/work/" + cur_proj.name + "+" + str(tran_task.task_id)
                    task_item_list.append(task_item)
            return jsonify(task_item_list)

        else: 
            return "fasle"
    else:
        return "false"

@app.route('/translator', methods=['POST','GET'])
def toTranslatorPage():
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role != 0:
            return "false"
        else:
            return render_template('translator_homepage.html', user=user1)
    else: 
        return "false"

@app.route('/work/<list:ids>', methods=['POST','GET'])
def work_zone(ids):
    # translation_path_file = '../static/uploads/'
    # translation_fname = secure_filename(filename)
    # log.info('translation_path_file = %s', translation_path_file)
    # log.info('translation_fname = %s', translation_fname)
    # return render_template('workingpage.html', translationId = translation_path_file + translation_fname)
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role is None:
            return "false"
        taskId = ids[1]
        cur_task = task.query.filter_by(id=taskId).first()
        source_path = cur_task.task_path
        target_path = source_path[0: source_path.rindex(".")] + ".txt"
        if not os.path.exists(source_path):
            return "No source file"
        if not os.path.exists(target_path):
            return "No target file"
        htmlContent = ""
        with open(target_path, 'r') as f:
            htmlContent = f.read()
        return render_template('workingpage.html', user=user1, taskId=taskId, translationId = source_path,
            htmlContent = htmlContent, targetPath=target_path)

@app.route('/save_text/<path:path>', methods=['POST','GET'])
def save_text(path):
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role != 0:
            return "false"
        if request.method == 'POST':
            html_content = json.loads(request.get_data())
            with open("./"+path, "w") as f:
                f.write(html_content)
            return "true"
        else:
            return "false"
    else: 
        return "false"

@app.route('/submit_text/<list:ids>', methods=['POST','GET'])
def submit_text(ids):
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role != 0:
            return "false"
        if request.method == 'POST':
            tr_ta = translators_tasks.query.filter_by(task_id=ids[0], translator_id=ids[1]).first()
            tr_ta.task_type = 'finished'
            db.session.add(tr_ta)
            db.session.commit()
            return "true"
        else:
            return "false"
    else: 
        return "false"

@app.route('/export/<name>', methods=['POST','GET'])
def export_project(name):
    if session.get('email'):
        email = session.get('email')
        user1 = user.query.filter_by(email=email).first()
        if user1.role != 1:
            return "false"
        # if request.method == 'POST':
        cur_proj = project.query.filter_by(name=name).first()
        tasks = task.query.filter_by(project_id=cur_proj.id).all()
        path_page = []
        for cur_task in tasks:
            path_page.append([cur_task.task_path, cur_task.start_page])
        
        path_page = sorted(path_page, key=lambda x: x[1])
        
        cur_file_name = name + '.docx'
        cur_file = './data/users/' + email + '/' + name +'/'
        cur_file = os.path.abspath(cur_file)
        if not os.path.exists(cur_file):
            os.makedirs(cur_file)
        myfile = b""
        document = Document()
        for path in path_page:
            txt_path = path[0][0: path[0].rindex(".")] + ".txt"
            with open(txt_path, "rb") as f:
                myfile += f.read()
        
        # myfile=myfile.decode("utf-8","xmlcharrefreplace")
        #myfile = re.sub(r'[^\x00-\x7F]+|\x0c',' ', myfile) # remove all non-XML-compatible characters
        myfile = myfile.decode("utf-8", "ignore")
        p = document.add_paragraph(myfile)
        document.save(os.path.join(cur_file, cur_file_name))
        response = make_response(send_from_directory(cur_file, cur_file_name, as_attachment=True))
        response.headers["Content-Disposition"] = "attachment; filename={}".format(cur_file_name.encode().decode('latin-1'))
        return response
    else:
        return "false"
